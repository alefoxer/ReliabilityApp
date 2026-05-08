from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ACCENT = RGBColor(0x0F, 0x5C, 0x6E)


def style_run(run, *, size=11, bold=False, color=None, italic=False):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color


def heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    style_run(run, size=14, bold=True, color=ACCENT)


def para(doc: Document, text: str, *, align=None, after=6) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    style_run(run)


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.1
        style_run(p.add_run(item))


def numbers(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.1
        style_run(p.add_run(item))


def mono_line(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    style_run(run, bold=True, color=ACCENT)


def build() -> Path:
    out_dir = Path("deliverables")
    out_dir.mkdir(exist_ok=True)
    target = out_dir / "pdf_snhe_verification_guide.docx"

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(1.8)
    sec.bottom_margin = Cm(1.8)
    sec.left_margin = Cm(2.2)
    sec.right_margin = Cm(2.2)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    style_run(title.add_run("Проверка примера из PDF\nдля СНЭ ЭМРТУ"), size=18, bold=True, color=ACCENT)

    para(
        doc,
        "Готовый документ для показа преподавателю: что именно загружать в программе, "
        "какие данные используются и какие результаты должны получиться.",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=14,
    )

    heading(doc, "1. Что именно проверяется")
    para(
        doc,
        "По PDF лучше всего проверяется верхнеуровневая структурная схема СНЭ как "
        "последовательное соединение шести подсистем. Это соответствует разделу 9.1.2 PDF "
        "и числам из таблиц 4 и 5.",
    )
    bullets(
        doc,
        [
            "Энергоблок МКТН.563255.008",
            "Энергоблок собственных нужд МКТН.563255.007",
            "Подсистема управления электропитанием МКТН.794775.025",
            "Комплект силовых кабельных линий МКТН.468921.082",
            "Комплект информационных кабельных линий МКТН.468921.083",
            "Распределительное устройство проходных коробок МКТН.674211.009",
        ],
    )

    heading(doc, "2. Что нужно сделать в программе")
    numbers(
        doc,
        [
            "Открыть модуль «Графический редактор».",
            "Нажать «Загрузить схему».",
            "Выбрать файл pdf_snhe_158h_scheme.json.",
            "В блоке «Параметры расчета схемы» установить метод «Аналитический расчёт».",
            "Установить горизонт t = 158.",
            "Нажать «Проверить схему».",
            "Нажать «Сгенерировать формулу».",
            "Нажать «Рассчитать схему».",
        ],
    )
    para(doc, "Ожидаемый смысл формулы:", after=4)
    mono_line(doc, "Pсист(t) = B1 · B2 · B3 · B4 · B5 · B6")
    mono_line(doc, "Kг_сист = K_B1 · K_B2 · K_B3 · K_B4 · K_B5 · K_B6")

    heading(doc, "3. Исходные данные из PDF")
    para(doc, "Данные по верхнему уровню СНЭ:", after=4)
    bullets(
        doc,
        [
            "B1: Энергоблок МКТН.563255.008. P(158 ч)=0.999988706, P(125 ч)=0.999999941, Tv=5 ч, Kg=0.999999997.",
            "B2: Энергоблок собственных нужд МКТН.563255.007. P(158 ч)=0.999918418, P(125 ч)=0.999999577, Tv=5 ч, Kg=0.999999980.",
            "B3: Подсистема управления электропитанием МКТН.794775.025. P(158 ч)=0.999248387, P(125 ч)=0.999996104, Tv=3 ч, Kg=0.999999907.",
            "B4: Комплект силовых кабельных линий МКТН.468921.082. P(158 ч)=0.999829933, P(125 ч)=0.999999119, Tv=6 ч, Kg=0.999999958.",
            "B5: Комплект информационных кабельных линий МКТН.468921.083. P(158 ч)=0.999914963, P(125 ч)=0.999999559, Tv=6 ч, Kg=0.999999979.",
            "B6: Распределительное устройство проходных коробок МКТН.674211.009. P(158 ч)=0.999999500, P(125 ч)=0.999999997, Tv=3 ч, Kg=0.999999980.",
        ],
    )

    heading(doc, "4. Какие файлы уже подготовлены")
    bullets(
        doc,
        [
            "pdf_snhe_158h_scheme.json: сценарий для воспроизведения таблицы 4 на t = 158 ч.",
            "pdf_snhe_125h_scheme.json: сценарий для воспроизведения таблицы 4 на t = 125 ч.",
            "Оба файла уже содержат параметры блоков и могут быть сразу загружены в графический редактор.",
        ],
    )
    para(doc, "Параметры, вшитые в файл pdf_snhe_158h_scheme.json:", after=4)
    bullets(
        doc,
        [
            "B1: lambda = 7.148E-08, Tv = 5, Kg = 0.999999997.",
            "B2: lambda = 5.164E-07, Tv = 5, Kg = 0.999999980.",
            "B3: lambda = 4.759E-06, Tv = 3, Kg = 0.999999907.",
            "B4: lambda = 1.076E-06, Tv = 6, Kg = 0.999999958.",
            "B5: lambda = 5.382E-07, Tv = 6, Kg = 0.999999979.",
            "B6: lambda = 3.165E-09, Tv = 3, Kg = 0.999999980.",
        ],
    )

    heading(doc, "5. Ожидаемые результаты")
    para(doc, "Сценарий A. Файл pdf_snhe_158h_scheme.json при t = 158 ч:", after=4)
    bullets(
        doc,
        [
            "P(B1)=0.999988706",
            "P(B2)=0.999918418",
            "P(B3)=0.999248387",
            "P(B4)=0.999829933",
            "P(B5)=0.999914963",
            "P(B6)=0.999999500",
            "Pсист(158 ч)=0.998900209",
            "Kг_сист=0.999999801",
        ],
    )
    para(doc, "Сценарий B. Файл pdf_snhe_125h_scheme.json при t = 125 ч:", after=4)
    bullets(
        doc,
        [
            "P(B1)=0.999999941",
            "P(B2)=0.999999577",
            "P(B3)=0.999996104",
            "P(B4)=0.999999119",
            "P(B5)=0.999999559",
            "P(B6)=0.999999997",
            "Pсист(125 ч)=0.999994298",
            "Kг_сист=0.999999801",
        ],
    )
    para(
        doc,
        "Важно: в самом PDF есть небольшая несостыковка по итоговому Kг. "
        "Произведение блочных значений из таблицы 5 дает 0.999999801, "
        "а в одной итоговой строке документа указано 0.99999982. "
        "Для защиты лучше показывать именно произведение значений таблицы 5.",
    )

    heading(doc, "6. Что программа пока не покрывает напрямую")
    bullets(
        doc,
        [
            "Раздел 9.2 PDF: функциональная надежность с резервом времени tз и tи.",
            "Раздел 11 PDF: долговечность, средний ресурс, срок службы, общий срок службы.",
            "Специальные формулы PDF по резервированию и восстановлению, которые не вынесены в отдельную методику интерфейса.",
        ],
    )
    para(
        doc,
        "Корректная формулировка для защиты: текущая версия программы полностью воспроизводит "
        "структурную проверку верхнеуровневой схемы СНЭ по вероятности безотказной работы и "
        "коэффициенту готовности, а специальные модели функциональной надежности и долговечности "
        "из данного PDF пока не вынесены в отдельные автоматические методики.",
    )

    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(footer.add_run("Подготовлено для проверки примера из PDF по СНЭ ЭМРТУ"), size=9, color=ACCENT)

    doc.save(target)
    print(target)
    return target


if __name__ == "__main__":
    build()
