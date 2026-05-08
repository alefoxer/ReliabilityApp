"""Professional report exporters for reliability calculations.

The UI prepares one ``ReportData`` object. This module normalizes that model
and renders it into format-specific outputs: HTML for local viewing/PDF,
DOCX for a readable engineering document, and XLSX for analytical workbooks.
"""

from __future__ import annotations

import io
import math
from pathlib import Path
from typing import Any

from PyQt6.QtCore import QSizeF
from PyQt6.QtGui import QPageSize, QTextDocument
from PyQt6.QtPrintSupport import QPrinter

from app.formulas.formula_rendering import (
    is_renderable_latex_formula,
    latex_to_html,
    readable_formula_text,
    render_latex_lines_to_png_bytes,
    render_latex_to_png_bytes,
    split_latex_formula_for_display,
)
from app.import_export.import_safety import safe_html as escape
from app.import_export.import_safety import safe_xlsx_value
from PIL import Image
from app.core.rbd_models import ReportData


APP_NAME = "Надёжность технических средств"
DEFAULT_METHOD = "Не указана"
DEFAULT_CALCULATION_METHOD = "Аналитический расчёт"
THRESHOLD_GREATER_IS_BETTER = {"P", "Kg", "Kog", "T0"}


def _threshold_passes(metric: str, value: float, threshold: float) -> bool:
    if metric in THRESHOLD_GREATER_IS_BETTER:
        return float(value) >= float(threshold)
    return float(value) <= float(threshold)


def _public_method_name(value: str) -> str:
    text = str(value or "").strip()
    if not text:
        return DEFAULT_METHOD
    if "AUTO.COMPOSITION" in text or text.startswith("Схема: AUTO."):
        return "Формула по структуре схемы"
    if text.startswith("AUTO."):
        return "Формула по структуре схемы"
    return text


def _public_formula_text(value: str) -> str:
    lines: list[str] = []
    for raw_line in str(value or "").splitlines():
        line = raw_line.strip()
        if not line:
            lines.append(raw_line)
            continue
        if "AUTO.COMPOSITION" in line or line.startswith("AUTO."):
            continue
        lines.append(raw_line)
    return "\n".join(lines).strip()


def _public_source_label(value: str) -> str:
    mapping = {
        "calculator": "Калькулятор",
        "editor": "Графический редактор",
        "import": "Импортированный расчёт",
    }
    text = str(value or "").strip()
    return mapping.get(text, text or "не указан")


def export_html(path: str | Path, report: ReportData) -> Path:
    target = Path(path)
    target.write_text(report_to_html(report), encoding="utf-8")
    return target


def export_txt(path: str | Path, report: ReportData) -> Path:
    target = Path(path)
    model = _report_model(report)
    lines = [
        model["project_name"],
        model["title"],
        model["subtitle"],
        f"Дата формирования: {model['created_at']}",
        f"Схема: {model['scheme_name']}",
        f"Методика: {model['method_name']}",
        f"Способ расчёта: {model['calculation_method']}",
        "",
        "ИТОГОВЫЙ ВЫВОД:",
        model["status_text"],
        "",
        "ИСХОДНЫЕ ДАННЫЕ:",
    ]
    lines.extend(f"- {name}: {value}" for name, value in model["inputs"])
    lines.extend(["", "РЕЗУЛЬТАТЫ:"])
    lines.extend(" | ".join(str(value) for value in row) for row in model["result_rows"])
    if model["formula_text"]:
        lines.extend(["", "ФОРМУЛЫ:", model["formula_text"]])
    if model["scheme_images"]:
        lines.extend(["", "СХЕМЫ СИСТЕМЫ:"])
        lines.extend(f"- {item['title']}: {item['path']}" for item in model["scheme_images"])
    if model["charts"]:
        lines.extend(["", "ГРАФИКИ:"])
        lines.extend(f"- {chart}" for chart in model["charts"])
    for title, rows in model["tables"].items():
        lines.extend(["", title])
        lines.extend(" | ".join(str(value) for value in row) for row in rows)
    if model["nomenclature_rows"]:
        lines.extend(["", "НОМЕНКЛАТУРА ПОКАЗАТЕЛЕЙ:"])
        lines.extend(f"- {label}: {value}" for label, value in model["nomenclature_rows"])
        if model["nomenclature_metrics"]:
            lines.append(f"- Рекомендуемые показатели: {model['nomenclature_metrics']}")
    if model["check_rows"]:
        lines.extend(["", "ПРОВЕРКИ:"])
        lines.extend(f"- {label}: {value}" for label, value in model["check_rows"])
    lines.extend(
        [
            "",
            "МЕТОДИКА И СЦЕНАРИЙ:",
            model["methodology"],
            "",
            "ЗАКЛЮЧЕНИЕ:",
            model["final_conclusion"],
        ]
    )
    if len(model["conclusion_rows"]) > 2:
        lines.extend(["", "ПРИМЕЧАНИЯ К ЗАКЛЮЧЕНИЮ:"])
        for label, value in model["conclusion_rows"][2:]:
            lines.append(f"- {label}: {value}")
    target.write_text("\n".join(lines), encoding="utf-8")
    return target


def export_pdf(path: str | Path, report: ReportData) -> Path:
    target = Path(path)
    document = QTextDocument()
    document.setHtml(report_to_html(report))
    printer = QPrinter(QPrinter.PrinterMode.HighResolution)
    printer.setOutputFormat(QPrinter.OutputFormat.PdfFormat)
    printer.setOutputFileName(str(target))
    printer.setPageSize(QPageSize(QPageSize.PageSizeId.A4))
    document.setPageSize(QSizeF(printer.pageRect(QPrinter.Unit.Point).size()))
    document.print(printer)
    return target


def report_to_html(report: ReportData) -> str:
    """Render a complete, local, self-contained HTML report shell."""
    model = _report_model(report)
    scheme_images = _html_scheme_block(model["scheme_images"])
    chart_images = "".join(_html_image_section("График результата", chart) for chart in model["charts"])
    graph_tables = _html_tables(model["tables"])
    formula_section = model["formula_html"] or f"<pre>{escape(model['formula_text'] or 'Формулы не указаны.')}</pre>"
    nomenclature_section = (
        f"<section><h2>Номенклатура показателей</h2>{_html_key_value_table(model['nomenclature_rows'])}"
        f"{_html_info_block('Рекомендуемые показатели', model['nomenclature_metrics'])}</section>"
        if model["nomenclature_rows"]
        else ""
    )

    return f"""<!DOCTYPE html>
<html lang="ru">
<head>
    <meta charset="utf-8">
    <title>{escape(model["title"])}</title>
    <style>
        body {{ font-family: "Segoe UI", sans-serif; margin: 34px; color: #172033; line-height: 1.45; }}
        .title {{ border-bottom: 3px solid #1f4f78; padding-bottom: 18px; margin-bottom: 22px; }}
        h1 {{ color: #12395c; margin: 0 0 8px 0; font-size: 28px; }}
        h2 {{ color: #1f4f78; margin-top: 24px; border-bottom: 1px solid #dbe5f1; padding-bottom: 5px; }}
        .meta {{ color: #536274; margin: 3px 0; }}
        .status {{ padding: 14px; border-radius: 10px; margin: 12px 0; font-weight: 700; }}
        .pass {{ background: #ecfdf3; border: 1px solid #86efac; color: #166534; }}
        .fail {{ background: #fff1f2; border: 1px solid #fda4af; color: #9f1239; }}
        .neutral {{ background: #f8fafc; border: 1px solid #cbd5e1; color: #334155; }}
        table {{ width: 100%; border-collapse: collapse; margin-top: 10px; }}
        td, th {{ border: 1px solid #cfd8e3; padding: 8px; vertical-align: top; }}
        th {{ background: #eef4fb; text-align: left; }}
        pre {{ white-space: pre-wrap; background: #f8fafc; border: 1px solid #cbd5e1; padding: 12px; border-radius: 8px; }}
        .report-image {{ max-width: 100%; border: 1px solid #d8e0ea; border-radius: 8px; padding: 6px; background: white; }}
    </style>
</head>
<body>
    <section class="title">
        <h1>{escape(model["project_name"])}</h1>
        <h2>{escape(model["title"])}</h2>
        <p class="meta">{escape(model["subtitle"])}</p>
        <p class="meta">Дата формирования: {escape(model["created_at"])}</p>
        <p class="meta">Схема: {escape(model["scheme_name"])}</p>
        <p class="meta">Методика: {escape(model["method_name"])}</p>
        <p class="meta">Способ расчёта: {escape(model["calculation_method"])}</p>
    </section>
    <section><h2>Итоговый вывод</h2><div class="status {model["status_class"]}">{escape(model["status_text"])}</div></section>
    <section><h2>Исходные данные</h2>{_html_key_value_table(model["inputs"])}</section>
    <section><h2>Числовые результаты</h2>{_html_result_table(model["result_rows"])}</section>
    {scheme_images}
    <section><h2>Методика и сценарий</h2><pre>{escape(model["methodology"])}</pre></section>
    {nomenclature_section}
    <section><h2>Формулы</h2>{formula_section}</section>
    {chart_images}
    {graph_tables}
    <section><h2>Проверки</h2>{_html_key_value_table(model["check_rows"])}</section>
    <section><h2>Заключение</h2><p>{escape(model["final_conclusion"])}</p>{_html_key_value_table(model["conclusion_rows"])}</section>
</body>
</html>"""


def export_docx(path: str | Path, report: ReportData) -> Path:
    """Export a formatted Word document suitable for submission or printing."""
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt
    except ImportError as exc:
        raise RuntimeError("Для экспорта DOCX установите зависимость python-docx.") from exc

    target = Path(path)
    model = _report_model(report)
    formula_assets_dir = target.parent
    document = Document()
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(model["project_name"])
    run.bold = True
    run.font.size = Pt(22)
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(model["title"]).bold = True
    document.add_paragraph(model["subtitle"])
    _add_docx_key_values(document, "Титульный блок", model["passport_rows"])
    _add_docx_status(document, model)
    _add_docx_table(document, "Краткая сводка", ["Параметр", "Значение"], model["summary_rows"])
    _add_docx_table(document, "Исходные данные", ["Параметр", "Значение"], model["inputs"])
    _add_docx_table(
        document,
        "Результаты расчёта",
        ["Показатель", "Обозначение", "Значение", "Единицы", "Комментарий"],
        model["result_rows"],
    )
    document.add_heading("Формулы и методика", level=1)
    document.add_paragraph(model["methodology_text"])
    document.add_heading("Формулы", level=1)
    _add_docx_formula_blocks(document, _formula_render_blocks(report), formula_assets_dir, Inches(6.1))
    if model["nomenclature_rows"]:
        _add_docx_table(document, "Номенклатура показателей", ["Параметр", "Выбор"], model["nomenclature_rows"])
        if model["nomenclature_metrics"]:
            document.add_paragraph("Рекомендуемые показатели: " + model["nomenclature_metrics"])
    document.add_heading("Графики и схема", level=1)
    for item in model["scheme_images"]:
        document.add_paragraph(str(item["title"]))
        _add_docx_image(document, "", str(item["path"]), Inches(6.1))
    for index, chart in enumerate(model["charts"], start=1):
        document.add_paragraph(f"График результата {index}")
        _add_docx_image(document, "", chart, Inches(6.1))
    for section in model["graph_sections"]:
        _add_docx_table(document, section["title"], section["headers"], section["rows"])
    _add_docx_table(document, "Проверки", ["Параметр", "Значение"], model["check_rows"])
    document.add_heading("Заключение", level=1)
    _add_docx_table(document, "Заключение", ["Параметр", "Значение"], model["conclusion_rows"])
    document.save(target)
    return target


def export_xlsx(path: str | Path, report: ReportData) -> Path:
    """Export an analytical workbook with summary, formulas, data and charts."""
    try:
        from openpyxl import Workbook
        from openpyxl.drawing.image import Image as XlsxImage
        from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
        from openpyxl.utils import get_column_letter
    except ImportError as exc:
        raise RuntimeError("Для экспорта XLSX установите зависимости openpyxl и Pillow.") from exc

    target = Path(path)
    model = _report_model(report)
    formula_assets_dir = target.parent
    wb = Workbook()

    summary = wb.active
    summary.title = "Сводка"
    _write_rows(summary, [[model["project_name"]], [model["title"]], ["Параметр", "Значение"], *model["summary_rows"]])
    _write_rows(
        wb.create_sheet("Паспорт отчёта"),
        [["Параметр", "Значение"], *model["passport_rows"], ["Навигация", "Исходные данные | Результаты | Формулы | Графики | Проверки | Заключение"]],
    )
    _write_rows(wb.create_sheet("Исходные данные"), [["Параметр", "Значение"], *model["inputs"]])
    _write_rows(
        wb.create_sheet("Результаты"),
        [["Показатель", "Обозначение", "Значение", "Единицы", "Комментарий"], *model["result_rows"]],
    )
    formulas_sheet = wb.create_sheet("Формулы")
    _write_rows(formulas_sheet, [["Раздел", "Содержание"]])
    _add_xlsx_formula_blocks(formulas_sheet, _formula_render_blocks(report), formula_assets_dir, XlsxImage)
    _write_rows(wb.create_sheet("Методика"), [["Раздел", "Содержание"], ["Методика", model["methodology_text"]]])
    nomenclature_rows = [["Параметр", "Выбор"], *model["nomenclature_rows"]]
    if model["nomenclature_metrics"]:
        nomenclature_rows.append(["Рекомендуемые показатели", model["nomenclature_metrics"]])
    _write_rows(wb.create_sheet("Номенклатура"), nomenclature_rows)

    graph_data = wb.create_sheet("Данные графиков")
    _write_rows(graph_data, model["graph_rows"])
    graph_row_offset = len(model["graph_rows"]) + 3
    for section in model["extra_table_sections"]:
        _write_rows(graph_data, [[section["title"]], section["headers"], *section["rows"]], start_row=graph_row_offset)
        graph_row_offset += len(section["rows"]) + 4

    charts = wb.create_sheet("Графики")
    if model["charts"]:
        _add_xlsx_image_with_reserved_space(charts, model["charts"][0], 2, 1, XlsxImage, max_width_px=820, max_height_px=420)
    else:
        _write_rows(charts, [["Раздел", "Содержание"], ["График", "Изображение графика не указано."]])

    scheme_sheet = wb.create_sheet("Схема")
    if not _add_xlsx_image_with_reserved_space(scheme_sheet, model["scheme_image_path"], 1, 1, XlsxImage, max_width_px=900, max_height_px=520):
        _write_rows(scheme_sheet, [["Раздел", "Содержание"], ["Схема", model["scheme_image_path"] or "Изображение схемы не указано."]])
    if len(model["scheme_images"]) > 1:
        subschemes_sheet = wb.create_sheet("Подсхемы")
        current_row = 1
        for item in model["scheme_images"][1:]:
            title = str(item.get("title", "Подсхема"))
            path_text = str(item.get("path", ""))
            subschemes_sheet.cell(row=current_row, column=1, value=safe_xlsx_value(title))
            next_row = _add_xlsx_image_with_reserved_space(subschemes_sheet, path_text, current_row + 1, 1, XlsxImage, max_width_px=900, max_height_px=520)
            if not next_row:
                subschemes_sheet.cell(
                    row=current_row + 1,
                    column=1,
                    value=safe_xlsx_value(path_text or "Изображение подсхемы недоступно."),
                )
                current_row += 3
            else:
                current_row = next_row + 1

    _write_rows(wb.create_sheet("Проверки"), [["Параметр", "Значение"], *model["check_rows"]])
    _write_rows(wb.create_sheet("Заключение"), [["Параметр", "Значение"], *model["conclusion_rows"]])

    _add_xlsx_title_sheet(wb, model)

    thin = Side(style="thin", color="CBD5E1")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    header_fill = PatternFill("solid", fgColor="DDEBF7")
    for sheet in wb.worksheets:
        for row in sheet.iter_rows():
            for cell in row:
                cell.alignment = Alignment(wrap_text=True, vertical="top")
                cell.border = border
                if cell.row == 1:
                    cell.fill = header_fill
                    cell.font = Font(bold=True, size=12)
                if str(cell.value).lower() in {"соответствует", "pass"}:
                    cell.fill = PatternFill("solid", fgColor="DCFCE7")
                if str(cell.value).lower() in {"не соответствует", "fail"}:
                    cell.fill = PatternFill("solid", fgColor="FEE2E2")
        for column in range(1, min(sheet.max_column, 8) + 1):
            sheet.column_dimensions[get_column_letter(column)].width = 24
        if sheet.max_row > 1:
            sheet.freeze_panes = "A2"
            if sheet.max_column >= 2:
                sheet.auto_filter.ref = sheet.dimensions
    wb.save(target)
    return target


def _report_model(report: ReportData) -> dict[str, Any]:
    calculation_method = (
        getattr(report, "calculation_method", "")
        or report.inputs.get("Способ расчёта")
        or report.inputs.get("Способ расчета")
        or DEFAULT_CALCULATION_METHOD
    )
    warnings = list(getattr(report, "warnings", []) or [])
    limitations = list(getattr(report, "limitations", []) or [])
    if report.notes:
        limitations.extend(line for line in str(report.notes).splitlines() if line.strip())
    metadata = getattr(report, "metadata", {}) or {}
    nomenclature = metadata.get("nomenclature", {}) if isinstance(metadata, dict) else {}
    nomenclature_rows = _nomenclature_rows(nomenclature)
    formula_package = getattr(report, "formula_package", None)
    formula_text = report.formula_text or report.formula_latex or ""
    formula_html = ""
    if formula_package is not None:
        formula_text = formula_package.plain_text or formula_package.latex_text
        formula_html = formula_package.html_text or ""
        warnings.extend(warning for warning in formula_package.warnings if warning not in warnings)
        if formula_package.limitations and formula_package.limitations not in limitations:
            limitations.append(formula_package.limitations)
    elif report.formula_latex:
        formula_html = latex_to_html(report.formula_latex)
    status_text = _status_text(report)
    key_result_rows = _key_result_rows(report.results)
    formula_rows = _formula_rows(report, formula_text)
    graph_rows = _graph_data_rows(report.tables)
    graph_sections, extra_table_sections = _split_table_sections(report.tables)
    check_rows = _check_rows(report, status_text)
    scheme_images = _normalize_scheme_images(report)
    summary_rows = [
        ["Статус", status_text],
        ["Схема", report.scheme_name or "не указана"],
        ["Методика", _public_method_name(report.method_name)],
        ["Способ расчёта", calculation_method],
        ["Пороговый показатель", report.threshold_metric or "не указан"],
        ["Порог", _format_value(report.threshold_value) if report.threshold_value is not None else "не задан"],
        *key_result_rows,
    ]
    passport_rows = [
        ["Дата формирования", report.created_at.strftime("%d.%m.%Y %H:%M")],
        ["Название отчёта", report.title or "Отчёт по расчёту надёжности"],
        ["Подзаголовок", report.subtitle or "Инженерный отчёт"],
        ["Схема", report.scheme_name or "не указана"],
        ["Методика", _public_method_name(report.method_name)],
        ["Способ расчёта", calculation_method],
        ["Источник данных", _public_source_label(str(metadata.get("source", "")))],
        ["Формула", "есть" if _public_formula_text(formula_text) else "нет"],
        ["График", "есть" if report.charts else "нет"],
        ["Схема-изображение", "есть" if scheme_images else "нет"],
    ]
    warning_rows = [[item] for item in warnings + limitations] or [["Не указаны."]]
    conclusion_rows = [
        ["Итоговый вывод", status_text],
        ["Заключение", report.final_conclusion or status_text],
    ]
    if warning_rows and warning_rows != [["Не указаны."]]:
        conclusion_rows.append(["Примечания", "\n".join(str(row[0]) for row in warning_rows)])
    return {
        "project_name": report.project_name or APP_NAME,
        "title": report.title or "Отчёт по расчёту надёжности",
        "subtitle": report.subtitle or "Инженерный отчёт",
        "created_at": report.created_at.strftime("%d.%m.%Y %H:%M"),
        "scheme_name": report.scheme_name or "не указана",
        "method_name": _public_method_name(report.method_name),
        "calculation_method": calculation_method,
        "inputs": _key_value_rows(report.inputs),
        "result_rows": _result_rows(report.results),
        "methodology": report.methodology or "Методика не указана.",
        "formula_text": _public_formula_text(formula_text),
        "formula_html": formula_html,
        "formula_rows": formula_rows,
        "scheme_image_path": scheme_images[0]["path"] if scheme_images else report.scheme_image_path,
        "scheme_images": scheme_images,
        "charts": [str(chart) for chart in report.charts if chart],
        "tables": report.tables,
        "graph_rows": graph_rows,
        "graph_sections": graph_sections,
        "extra_table_sections": extra_table_sections,
        "threshold_metric": report.threshold_metric,
        "threshold_value": report.threshold_value,
        "threshold_passed": report.threshold_passed,
        "status_text": status_text,
        "status_class": _status_class(report),
        "summary_rows": summary_rows,
        "passport_rows": passport_rows,
        "check_rows": check_rows,
        "warning_rows": warning_rows,
        "methodology_text": report.methodology or "Методика не указана.",
        "warnings": warnings,
        "limitations": limitations,
        "notes": report.notes,
        "metadata": metadata,
        "nomenclature_rows": nomenclature_rows,
        "nomenclature_metrics": str(nomenclature.get("recommended_metrics_text", "")).strip(),
        "conclusion_rows": conclusion_rows,
        "final_conclusion": report.final_conclusion or status_text,
    }


def _status_text(report: ReportData) -> str:
    if report.threshold_conclusion:
        return report.threshold_conclusion
    if report.threshold_value is not None:
        value = report.results.get(report.threshold_metric)
        if isinstance(value, (int, float)):
            metric = str(report.threshold_metric or "")
            passed = _threshold_passes(metric, float(value), float(report.threshold_value))
            verdict = "соответствует" if passed else "не соответствует"
            return (
                f"Система {verdict} заданному порогу: "
                f"{metric or 'показатель'}={float(value):.6f}, порог={float(report.threshold_value):.6f}."
            )
        return "Порог задан, но выбранный показатель отсутствует в результатах."
    return report.final_conclusion or "Расчёт выполнен. Итоговые показатели приведены в отчёте."


def _status_class(report: ReportData) -> str:
    if report.threshold_passed is True:
        return "pass"
    if report.threshold_passed is False:
        return "fail"
    return "neutral"


def _key_value_rows(values: dict[str, Any]) -> list[list[Any]]:
    return [[key, value] for key, value in values.items()]


def _result_rows(results: dict[str, Any]) -> list[list[Any]]:
    meta = {
        "P": ("Вероятность безотказной работы", "P(t)", "-", "Основной показатель надёжности за заданное время"),
        "Kg": ("Коэффициент готовности", "Kг", "-", "Доля времени, когда система работоспособна"),
        "Kog": ("Коэффициент оперативной готовности", "Kог", "-", "Готовность с учётом безотказной работы за интервал"),
        "T0": ("Средняя наработка до отказа", "T0", "ч", "Среднее время до отказа"),
        "Tv": ("Среднее время восстановления", "Tв", "ч", "Среднее время восстановления после отказа"),
        "Tpr": ("Среднее время простоя", "Tпр", "ч", "Средняя длительность простоя"),
    }
    rows: list[list[Any]] = []
    for key, value in results.items():
        name, symbol, unit, comment = meta.get(str(key), (str(key), str(key), "-", "Расчётный показатель"))
        rows.append([name, symbol, _format_value(value), unit, comment])
    return rows


def _format_value(value: Any) -> Any:
    if isinstance(value, float):
        return f"{value:.8g}"
    return value


def _html_key_value_table(rows: list[list[Any]]) -> str:
    body = "".join(f"<tr><td><b>{escape(str(k))}</b></td><td>{escape(str(v))}</td></tr>" for k, v in rows)
    return f"<table>{body}</table>"


def _html_result_table(rows: list[list[Any]]) -> str:
    headers = ["Показатель", "Обозначение", "Значение", "Единицы", "Комментарий"]
    return _html_table(headers, rows)


def _html_tables(tables: dict[str, list[tuple[Any, ...]]]) -> str:
    blocks: list[str] = []
    for title, rows in tables.items():
        row_lists = [list(row) for row in rows]
        blocks.append(f"<section><h2>{escape(str(title))}</h2>{_html_table(_table_headers(row_lists), row_lists)}</section>")
    return "".join(blocks)


def _html_table(headers: list[str], rows: list[list[Any]]) -> str:
    header_html = "".join(f"<th>{escape(str(header))}</th>" for header in headers)
    body = "".join(
        "<tr>" + "".join(f"<td>{escape(str(value))}</td>" for value in row) + "</tr>"
        for row in rows
    )
    return f"<table><thead><tr>{header_html}</tr></thead><tbody>{body}</tbody></table>"


def _html_image_section(title: str, path: str) -> str:
    if not path:
        return ""
    return f"<section><h2>{escape(title)}</h2><img src='{escape(path)}' class='report-image'></section>"


def _html_image_sections(items: list[dict[str, Any]]) -> str:
    if not items:
        return ""
    return "".join(_html_image_section(str(item.get("title", "Схема системы")), str(item.get("path", ""))) for item in items)


def _html_scheme_block(items: list[dict[str, Any]]) -> str:
    if not items:
        return ""
    parts = ["<section><h2>Схема системы</h2>"]
    for index, item in enumerate(items):
        path = str(item.get("path", ""))
        if not path:
            continue
        title = str(item.get("title", "")).strip()
        if index == 0:
            parts.append(f"<img src='{escape(path)}' class='report-image'>")
            continue
        if title:
            parts.append(f"<h3>{escape(title)}</h3>")
        parts.append(f"<img src='{escape(path)}' class='report-image'>")
    parts.append("</section>")
    return "".join(parts)


def _html_info_block(title: str, text: str) -> str:
    if not text:
        return ""
    return f"<p><b>{escape(title)}:</b> {escape(text)}</p>"


def _normalize_scheme_images(report: ReportData) -> list[dict[str, Any]]:
    raw_items = list(getattr(report, "scheme_images", []) or [])
    items: list[dict[str, Any]] = []
    for index, raw_item in enumerate(raw_items):
        if not isinstance(raw_item, dict):
            continue
        path = str(raw_item.get("path", "")).strip()
        if not path:
            continue
        title = str(raw_item.get("title", "")).strip()
        block_name = str(raw_item.get("block_name", "")).strip()
        level_value = raw_item.get("level", 0)
        try:
            level = int(level_value)
        except (TypeError, ValueError):
            level = 0
        if not title:
            title = "Схема системы" if index == 0 else f"Подсхема блока {block_name or index}"
        items.append({"title": title, "path": path, "level": level, "block_name": block_name})
    if items:
        return items
    fallback_path = str(getattr(report, "scheme_image_path", "") or "").strip()
    if not fallback_path:
        return []
    return [{"title": "Схема системы", "path": fallback_path, "level": 0, "block_name": ""}]


def _nomenclature_rows(nomenclature: dict[str, Any]) -> list[list[Any]]:
    if not nomenclature:
        return []
    rows: list[list[Any]] = []
    mapping = [
        ("Назначение", "purpose_label"),
        ("Режим применения", "usage_mode_label"),
        ("Восстановление", "recovery_mode_label"),
    ]
    for label, key in mapping:
        value = str(nomenclature.get(key, "")).strip()
        if value:
            rows.append([label, value])
    if "requires_tto" in nomenclature:
        rows.append(["Требуется Tто", "Да" if nomenclature.get("requires_tto") else "Нет"])
    tto_value = str(nomenclature.get("tto", "")).strip()
    if tto_value:
        rows.append(["Tто", tto_value])
    return rows


def _key_result_rows(results: dict[str, Any]) -> list[list[Any]]:
    preferred = [("P", "P(t)"), ("Kg", "Kг"), ("T0", "T0")]
    rows: list[list[Any]] = []
    for key, label in preferred:
        if key in results:
            rows.append([label, _format_value(results[key])])
    return rows


def _formula_rows(report: ReportData, formula_text: str) -> list[list[Any]]:
    rows = [["Итоговая формула", _public_formula_text(formula_text) or "Формулы не указаны."]]
    for key, value in report.results.items():
        if key in {"P", "Kg", "Kog", "T0", "Tv", "Tpr"}:
            rows.append([f"Формула показателя {key}", f"{key} = {_format_value(value)}"])
    return rows


def _formula_render_blocks(report: ReportData) -> list[dict[str, str]]:
    package = getattr(report, "formula_package", None)
    blocks: list[dict[str, str]] = []
    if package is not None:
        grouped_items = [
            ("Формулы", list(getattr(package, "formulas", []) or [])),
            ("Промежуточные формулы", list(getattr(package, "intermediate_formulas", []) or [])),
            ("Формулы показателей", list(getattr(package, "result_formulas", []) or [])),
        ]
        for section_label, items in grouped_items:
            for item in items:
                label = str(getattr(item, "label", "") or "").strip() or section_label
                latex_value = str(
                    getattr(item, "instantiated_latex", "")
                    or getattr(item, "display_latex", "")
                    or getattr(item, "general_latex", "")
                    or ""
                ).strip()
                plain_value = str(
                    getattr(item, "plain_text", "")
                    or getattr(item, "instantiated_formula", "")
                    or getattr(item, "symbolic_template", "")
                    or ""
                ).strip()
                if latex_value or plain_value:
                    blocks.append({"label": label, "latex": latex_value, "plain": plain_value})
        if not blocks and getattr(package, "plain_text", "").strip():
            blocks.append(
                {
                    "label": "Итоговая формула",
                    "latex": "",
                    "plain": str(package.plain_text or "").strip(),
                }
            )
    if blocks:
        return blocks
    fallback_plain = _public_formula_text(report.formula_text or report.formula_latex or "")
    if not fallback_plain:
        fallback_plain = "Формулы не указаны."
    fallback_latex = str(report.formula_latex or "").strip()
    return [{"label": "Итоговая формула", "latex": fallback_latex, "plain": fallback_plain}]


def _latex_formula_lines(value: str) -> list[str]:
    """Return only math lines that are safe to render as formula images."""
    lines: list[str] = []
    for raw_line in str(value or "").replace("\\[", "\n").replace("\\]", "\n").splitlines():
        line = raw_line.strip()
        if not line or not is_renderable_latex_formula(line):
            continue
        lines.extend(split_latex_formula_for_display(line, max_line_length=90))
    return [line for line in lines if line.strip()]


def _formula_png_bytes(block: dict[str, str]) -> bytes:
    latex_lines = _latex_formula_lines(str(block.get("latex", "") or ""))
    if latex_lines:
        try:
            png_data = render_latex_lines_to_png_bytes(latex_lines, font_size=15)
        except Exception:
            png_data = b""
        if png_data:
            return png_data
        if len(latex_lines) == 1:
            try:
                png_data = render_latex_to_png_bytes(latex_lines[0], font_size=15)
            except Exception:
                png_data = b""
            if png_data:
                return png_data
    return b""


def _xlsx_formula_png_bytes(block: dict[str, str]) -> bytes:
    png_data = _formula_png_bytes(block)
    if not png_data:
        return b""
    try:
        with Image.open(io.BytesIO(png_data)) as image:
            flattened = Image.new("RGB", image.size, (255, 255, 255))
            flattened.paste(image, mask=image.getchannel("A") if "A" in image.getbands() else None)
            buffer = io.BytesIO()
            flattened.save(buffer, format="PNG")
            return buffer.getvalue()
    except Exception:
        return png_data


def _formula_fallback_text(block: dict[str, str]) -> str:
    plain_value = str(block.get("plain", "") or "").strip()
    if plain_value:
        return "\n".join(readable_formula_text(line) for line in plain_value.splitlines() if line.strip()).strip()
    latex_value = str(block.get("latex", "") or "").strip()
    if latex_value:
        return readable_formula_text(latex_value).strip()
    return "Формула недоступна."


def _check_rows(report: ReportData, status_text: str) -> list[list[Any]]:
    metric_name = report.threshold_metric or "не указан"
    actual_value = report.results.get(report.threshold_metric)
    return [
        ["Пороговый показатель", metric_name],
        ["Порог", _format_value(report.threshold_value) if report.threshold_value is not None else "не задан"],
        ["Фактическое значение", _format_value(actual_value) if actual_value is not None else "отсутствует"],
        ["Статус", "соответствует" if report.threshold_passed is True else ("не соответствует" if report.threshold_passed is False else "не определён")],
        ["Пояснение", status_text],
    ]


def _split_table_sections(tables: dict[str, list[tuple[Any, ...]]]) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    graph_sections: list[dict[str, Any]] = []
    extra_sections: list[dict[str, Any]] = []
    for title, rows in tables.items():
        row_lists = [list(row) for row in rows]
        section = {
            "title": str(title),
            "headers": _table_headers(row_lists),
            "rows": row_lists,
        }
        lowered_title = str(title).lower()
        if "график" in lowered_title or "время" in lowered_title:
            graph_sections.append(section)
        else:
            extra_sections.append(section)
    return graph_sections, extra_sections


def _add_docx_key_values(document: Any, title: str, rows: list[tuple[Any, Any]]) -> None:
    _add_docx_table(document, title, ["Параметр", "Значение"], rows)


def _add_docx_status(document: Any, model: dict[str, Any]) -> None:
    document.add_heading("Итоговый вывод", level=1)
    paragraph = document.add_paragraph()
    run = paragraph.add_run(model["status_text"])
    run.bold = True


def _add_docx_table(document: Any, title: str, headers: list[str], rows: list[list[Any]] | list[tuple[Any, ...]]) -> None:
    if title:
        document.add_heading(title, level=1)
    table = document.add_table(rows=1, cols=max(1, len(headers)))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = str(header)
    for row in rows:
        values = list(row)
        cells = table.add_row().cells
        for index in range(len(headers)):
            cells[index].text = str(values[index]) if index < len(values) else ""


def _add_docx_formula_blocks(document: Any, blocks: list[dict[str, str]], assets_dir: Path, width: Any) -> None:
    for index, block in enumerate(blocks, start=1):
        label = str(block.get("label", "") or f"Формула {index}")
        document.add_paragraph(label)
        png_data = _formula_png_bytes(block)
        if png_data:
            width_px, height_px = _image_size_from_bytes(png_data)
            document.add_picture(
                io.BytesIO(png_data),
                width=_docx_image_width_from_size(width_px, height_px, max_width_inches=6.1, max_height_px=360, max_upscale=1.0),
            )
            continue
        document.add_paragraph(_formula_fallback_text(block))


def _add_docx_image(document: Any, title: str, path: str, width: Any) -> None:
    if not path or not Path(path).exists():
        return
    if title:
        document.add_heading(title, level=1)
    document.add_picture(path, width=_docx_image_width(path, max_width_inches=6.1, max_height_px=520, max_upscale=1.05))


def _image_size_px(path: str | Path) -> tuple[int, int]:
    try:
        with Image.open(path) as image:
            return image.size
    except Exception:
        return (1, 1)


def _image_size_from_bytes(data: bytes) -> tuple[int, int]:
    try:
        with Image.open(io.BytesIO(data)) as image:
            return image.size
    except Exception:
        return (1, 1)


def _normalize_image_dimensions(
    width_px: int,
    height_px: int,
    *,
    max_width_px: int,
    max_height_px: int,
    max_upscale: float = 1.0,
) -> tuple[int, int]:
    width_px = max(1, int(width_px or 1))
    height_px = max(1, int(height_px or 1))
    scale = min(max_width_px / width_px, max_height_px / height_px, max_upscale)
    scale = max(0.01, scale)
    return (max(1, int(width_px * scale)), max(1, int(height_px * scale)))


def _docx_image_width(path: str | Path, *, max_width_inches: float, max_height_px: int, max_upscale: float) -> Any:
    width_px, height_px = _image_size_px(path)
    return _docx_image_width_from_size(width_px, height_px, max_width_inches=max_width_inches, max_height_px=max_height_px, max_upscale=max_upscale)


def _docx_image_width_from_size(width_px: int, height_px: int, *, max_width_inches: float, max_height_px: int, max_upscale: float) -> Any:
    from docx.shared import Inches

    display_width_px, _ = _normalize_image_dimensions(
        width_px,
        height_px,
        max_width_px=int(max_width_inches * 96),
        max_height_px=max_height_px,
        max_upscale=max_upscale,
    )
    return Inches(display_width_px / 96)


def _write_rows(ws: Any, rows: list[list[Any]], start_row: int | None = None) -> None:
    if start_row is None:
        for row in rows:
            ws.append([safe_xlsx_value(value) for value in row])
        return
    current_row = start_row
    for row in rows:
        for column_index, value in enumerate(row, start=1):
            ws.cell(row=current_row, column=column_index, value=safe_xlsx_value(value))
        current_row += 1


def _add_xlsx_image(ws: Any, path: str, cell: str, image_cls: Any) -> bool:
    if not path or not Path(path).exists():
        return False
    try:
        ws.add_image(image_cls(path), cell)
        return True
    except Exception:
        return False


def _add_xlsx_image_with_reserved_space(
    ws: Any,
    path: str,
    row: int,
    col: int,
    image_cls: Any,
    *,
    max_width_px: int = 760,
    max_height_px: int = 360,
    max_upscale: float = 1.0,
    padding_rows: int = 1,
) -> int:
    if not path or not Path(path).exists():
        return 0
    try:
        from openpyxl.utils import get_column_letter

        width_px, height_px = _image_size_px(path)
        display_width_px, display_height_px = _normalize_image_dimensions(
            width_px,
            height_px,
            max_width_px=max_width_px,
            max_height_px=max_height_px,
            max_upscale=max_upscale,
        )
        image = image_cls(path)
        image.width = display_width_px
        image.height = display_height_px
        ws.add_image(image, f"{get_column_letter(col)}{row}")

        row_count = max(1, int(math.ceil((display_height_px + 10) / 20)))
        row_height_pt = max(15.0, (display_height_px * 0.75 + 8) / row_count)
        for row_index in range(row, row + row_count):
            current_height = ws.row_dimensions[row_index].height or 0
            ws.row_dimensions[row_index].height = max(current_height, row_height_pt)
        return row + row_count + max(0, padding_rows)
    except Exception:
        return 0


def _reserve_xlsx_image_rows(ws: Any, row: int, display_height_px: int, padding_rows: int) -> int:
    row_count = max(1, int(math.ceil((display_height_px + 10) / 20)))
    row_height_pt = max(15.0, (display_height_px * 0.75 + 8) / row_count)
    for row_index in range(row, row + row_count):
        current_height = ws.row_dimensions[row_index].height or 0
        ws.row_dimensions[row_index].height = max(current_height, row_height_pt)
    return row + row_count + max(0, padding_rows)


def _add_xlsx_png_bytes_with_reserved_space(
    ws: Any,
    png_data: bytes,
    row: int,
    col: int,
    image_cls: Any,
    *,
    max_width_px: int = 760,
    max_height_px: int = 360,
    max_upscale: float = 1.0,
    padding_rows: int = 1,
) -> int:
    if not png_data:
        return 0
    try:
        from openpyxl.utils import get_column_letter

        pil_image = Image.open(io.BytesIO(png_data))
        display_width_px, display_height_px = _normalize_image_dimensions(
            pil_image.width,
            pil_image.height,
            max_width_px=max_width_px,
            max_height_px=max_height_px,
            max_upscale=max_upscale,
        )
        image = image_cls(pil_image)
        image.width = display_width_px
        image.height = display_height_px
        ws.add_image(image, f"{get_column_letter(col)}{row}")
        return _reserve_xlsx_image_rows(ws, row, display_height_px, padding_rows)
    except Exception:
        return 0


def _add_xlsx_formula_blocks(ws: Any, blocks: list[dict[str, str]], assets_dir: Path, image_cls: Any) -> None:
    current_row = 2
    for index, block in enumerate(blocks, start=1):
        label = str(block.get("label", "") or f"Формула {index}")
        ws.cell(row=current_row, column=1, value=safe_xlsx_value(label))
        png_data = _xlsx_formula_png_bytes(block)
        if png_data:
            next_row = _add_xlsx_png_bytes_with_reserved_space(
                ws,
                png_data,
                current_row,
                2,
                image_cls,
                max_width_px=760,
                max_height_px=240,
                max_upscale=1.0,
            )
            if next_row:
                current_row = next_row
                continue
            else:
                ws.cell(row=current_row, column=2, value=safe_xlsx_value(_formula_fallback_text(block)))
        else:
            ws.cell(row=current_row, column=2, value=safe_xlsx_value(_formula_fallback_text(block)))
        fallback_lines = max(1, len(_formula_fallback_text(block).splitlines()))
        current_row += max(3, fallback_lines + 1)


def _add_xlsx_title_sheet(wb: Any, model: dict[str, Any]) -> None:
    ws = wb.create_sheet("Титульный лист", 0)
    _write_rows(
        ws,
        [
            [model["title"]],
            ["Дата формирования", model["created_at"]],
            ["Схема", model["scheme_name"]],
            ["Методика", model["method_name"]],
            ["Способ расчёта", model["calculation_method"]],
            ["Итоговый статус", model["status_text"]],
            [],
            ["Структура книги"],
            ["Лист", "Описание", "Ссылка"],
        ],
    )
    descriptions = {
        "Сводка": "Основные параметры и итоговый вывод",
        "Паспорт отчёта": "Паспорт отчёта и навигация",
        "Исходные данные": "Параметры расчёта и входные данные",
        "Результаты": "Таблица рассчитанных показателей",
        "Формулы": "Формулы и методика расчёта",
        "Методика": "Описание методики расчёта",
        "Номенклатура": "Номенклатура показателей",
        "Данные графиков": "Данные для построения графиков",
        "Графики": "Графики зависимости показателей",
        "Схема": "Изображение схемы системы",
        "Подсхемы": "Изображения вложенных схем",
        "Проверки": "Проверка порогов и статуса",
        "Заключение": "Итоговое заключение",
    }
    current_row = 10
    for sheet in wb.worksheets[1:]:
        ws.cell(row=current_row, column=1, value=safe_xlsx_value(sheet.title))
        ws.cell(row=current_row, column=2, value=safe_xlsx_value(descriptions.get(sheet.title, "Раздел отчёта")))
        link_cell = ws.cell(row=current_row, column=3, value="Перейти")
        link_cell.hyperlink = f"#'{sheet.title}'!A1"
        link_cell.style = "Hyperlink"
        back_cell = sheet.cell(row=1, column=8, value="На титульный лист")
        back_cell.hyperlink = "#'Титульный лист'!A1"
        back_cell.style = "Hyperlink"
        current_row += 1
    ws.column_dimensions["A"].width = 28
    ws.column_dimensions["B"].width = 54
    ws.column_dimensions["C"].width = 16
    ws.freeze_panes = "A9"


def _graph_data_rows(tables: dict[str, list[tuple[Any, ...]]]) -> list[list[Any]]:
    if not tables:
        return [["t", "P(t)"], ["Нет данных", ""]]
    first_rows = next(iter(tables.values()))
    rows = [list(row) for row in first_rows]
    if not rows:
        return [["t", "P(t)"], ["Нет данных", ""]]
    first = [str(value).lower() for value in rows[0]]
    if "t" not in first[0] and "время" not in first[0]:
        rows.insert(0, ["t", "P(t)"])
    return rows


def _table_headers(rows: list[Any]) -> list[str]:
    max_len = max((len(row) for row in rows), default=2)
    if max_len == 2:
        return ["Параметр", "Значение"]
    if max_len == 5:
        return ["Показатель", "Обозначение", "Значение", "Единицы", "Комментарий"]
    return [f"Колонка {index}" for index in range(1, max_len + 1)]
